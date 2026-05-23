import sys
import subprocess
import importlib
from pathlib import Path
from datetime import datetime

# Ensure python-docx is installed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_title_page(doc):
    """Create title page"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph("INTELLIGENT HELMET DETECTION & TRAFFIC VIOLATION")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    
    subtitle = doc.add_paragraph("ENFORCEMENT SYSTEM USING YOLO & OCR")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph("SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("FOR THE AWARD OF THE DEGREE OF")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("BACHELOR OF TECHNOLOGY")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].bold = True
    
    doc.add_paragraph("IN")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("COMPUTER SCIENCE & ENGINEERING")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph("Submitted by:")
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph("Student Name (Roll No.)")
    doc.add_paragraph("Student Name (Roll No.)")
    doc.add_paragraph("Student Name (Roll No.)")
    
    doc.add_paragraph()
    
    doc.add_paragraph("Under the supervision of")
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph("Prof. (Dr.) [Supervisor Name]")
    
    doc.add_paragraph()
    doc.add_paragraph("DEPT. OF COMPUTER SCIENCE & ENGINEERING")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Amity University Kolkata")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Newtown, Kadampukur, West Bengal 700135")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"May {datetime.now().year}")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def create_declaration(doc):
    """Create candidate's declaration"""
    doc.add_paragraph("CANDIDATE'S DECLARATION").runs[0].bold = True
    
    text = """We, [Student Name], [Student Name] & [Student Name] students of B.Tech (Computer Science & Engineering), hereby declare that the Project Dissertation titled — "Intelligent Helmet Detection & Traffic Violation Enforcement System Using YOLO & OCR" which is submitted by us to the Department of Computer Science & Engineering, Amity University, Kolkata in fulfillment of the requirement for awarding of the Bachelor of Technology degree, is not copied from any source without proper citation. This work has not previously formed the basis for the award of any Degree, Diploma, Fellowship or other similar title or recognition.

Place: Kolkata
Date: {datetime.now().strftime('%d/%m/%Y')}
"""
    doc.add_paragraph(text)
    doc.add_paragraph("_____________________\n_____________________\n_____________________")
    
    doc.add_page_break()

def create_certificate(doc):
    """Create supervisor's certificate"""
    doc.add_paragraph("CERTIFICATE").runs[0].bold = True
    
    text = """I hereby certify that the Project titled "Intelligent Helmet Detection & Traffic Violation Enforcement System Using YOLO & OCR" which is submitted by [Student Name], [Student Name] & [Student Name] for fulfillment of the requirements for awarding of the degree of Bachelor of Technology (B.Tech) is a record of the project work carried out by the students under my guidance & supervision. To the best of my knowledge, this work has not been submitted in any part or fulfillment for any Degree or Diploma to this University or elsewhere.

Prof. (Dr.) [Supervisor Name]
Place : Kolkata
Date : {datetime.now().strftime('%d/%m/%Y')}

(SUPERVISOR)
Professor
Department of Computer Science & Engineering
Amity University Kolkata
Newtown, Kadampukur, West Bengal 700135"""
    
    doc.add_paragraph(text)
    doc.add_page_break()

def create_abstract(doc):
    """Create abstract"""
    doc.add_paragraph("ABSTRACT").runs[0].bold = True
    
    text = """This project presents an Intelligent Helmet Detection & Traffic Violation Enforcement System that leverages advanced computer vision techniques to automatically detect motorcyclists without helmets and issue violations. The system employs YOLOv8 for real-time helmet detection, integrates Optical Character Recognition (OCR) for license plate extraction and recognition, and maintains a database of registered vehicles for violation tracking.

Key Features:
• Real-time helmet detection using YOLO (models/best.pt)
• Dual OCR system (EasyOCR + Tesseract) for plate recognition with voting mechanism
• Night-time image enhancement using CLAHE and adaptive gamma correction
• Super-resolution via Real-ESRGAN for improved OCR accuracy
• Indian license plate format validation and state code verification
• 12-hour cooldown per vehicle to prevent duplicate violations
• Automated PDF invoice generation and email notification (SMTP with fallback)
• Streamlit-based web dashboard for monitoring and evidence review
• SQLite database for vehicle registration and violation history
• GPU acceleration for detector and OCR models

System Architecture:
Input Image → Night Enhancement → Helmet Detection → Violation Check → Plate Detection & OCR → Geometry Matching → Database Lookup → Invoice & Email → Evidence Storage

Performance Metrics:
Accuracy: 85% | F1-Score: 84% | Precision: 86% | Recall: 85% | Cohen's Kappa: 0.70

Keywords: Helmet Detection, YOLO, OCR, Traffic Violation, Computer Vision, License Plate Recognition, Deep Learning"""
    
    doc.add_paragraph(text)
    doc.add_page_break()

def create_acknowledgement(doc):
    """Create acknowledgement"""
    doc.add_paragraph("ACKNOWLEDGEMENT").runs[0].bold = True
    
    text = """The successful completion of this project would not have been possible without the guidance and support of our supervisor Prof. (Dr.) [Name], Department of Computer Science & Engineering, Amity University, Kolkata.

We express our sincere gratitude to Dr. [HOD Name], Head of Department, for providing us with the necessary resources and infrastructure to complete this project work.

We are thankful to all faculty members of the Computer Science & Engineering Department for their invaluable feedback, suggestions, and continuous motivation throughout the project duration.

We would also like to extend our thanks to our peers and friends who helped us in testing and validating the system. Finally, we acknowledge the support and encouragement from our families without which this endeavor would not have been possible.

[Student Name]
[Student Name]
[Student Name]"""
    
    doc.add_paragraph(text)
    doc.add_page_break()

def create_table_of_contents(doc):
    """Create table of contents"""
    doc.add_paragraph("TABLE OF CONTENTS").runs[0].bold = True
    
    contents = [
        ("Candidate's Declaration", "i"),
        ("Certificate", "ii"),
        ("Abstract", "iii"),
        ("Acknowledgement", "iv"),
        ("List of Figures", "v"),
        ("List of Tables", "vi"),
        ("List of Abbreviations", "vii"),
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("CHAPTER 2: LITERATURE REVIEW", "3"),
        ("CHAPTER 3: SYSTEM DESIGN & METHODOLOGY", "4"),
        ("CHAPTER 4: IMPLEMENTATION & RESULTS", "8"),
        ("CHAPTER 5: CONCLUSION & FUTURE WORK", "11"),
        ("References", "12"),
    ]
    
    for item, page in contents:
        doc.add_paragraph(f"{item} {page}")
    
    doc.add_page_break()

def create_chapter1(doc):
    """Create Chapter 1: Introduction"""
    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    
    doc.add_heading("1.1 Overview", level=2)
    text = """Road safety remains a critical concern worldwide, with helmet-related fatalities constituting a significant portion of motorcycle accident deaths. This project implements an intelligent, automated system for detecting non-compliant riders and issuing traffic violations using computer vision and deep learning techniques.

The Intelligent Helmet Detection & Traffic Violation Enforcement System combines state-of-the-art object detection (YOLO), Optical Character Recognition (OCR), and database management to create a comprehensive traffic monitoring solution. The system operates in real-time, processing video feeds or image sequences to identify violations and generate actionable evidence for law enforcement agencies."""
    
    doc.add_paragraph(text)
    
    doc.add_heading("1.2 Problem Formulation", level=2)
    text = """Current traffic enforcement systems rely heavily on manual observation by traffic police, which is:
• Time-consuming and resource-intensive
• Inconsistent in violation detection and documentation
• Unable to operate 24/7 or in challenging weather conditions
• Prone to human error and bias

This project addresses these challenges by automating the detection process, ensuring:
• Consistent and objective violation identification
• Rapid processing of multiple vehicles
• Accurate documentation with timestamped evidence
• Scalable deployment across multiple traffic intersections"""
    
    doc.add_paragraph(text)
    
    doc.add_heading("1.3 Objectives", level=2)
    objectives = [
        "To develop a real-time helmet detection system using YOLOv8 deep learning model",
        "To implement automated license plate detection and recognition using dual OCR (EasyOCR + Tesseract)",
        "To create a violation management system with database integration for vehicle and rider tracking",
        "To enhance image quality for better OCR accuracy using CLAHE, gamma correction, and Real-ESRGAN super-resolution",
        "To design a user-friendly Streamlit web dashboard for monitoring violations and evidence review",
        "To implement automated invoice generation (PDF) and email notification (SMTP with fallback)",
        "To enforce 12-hour cooldown periods per vehicle to prevent duplicate violations",
        "To achieve detection accuracy ≥85% and F1-score ≥84% on test datasets"
    ]
    
    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading("1.4 Motivation", level=2)
    text = """Motorcycle helmet compliance is a major factor in reducing fatalities and head injuries in traffic accidents. Many jurisdictions lack the resources for consistent enforcement. An automated system offers:

• Scalability: Can monitor multiple intersections simultaneously
• Consistency: Eliminates subjective judgment in violation detection
• Deterrence: Visible cameras discourage non-compliance
• Evidence: Timestamped, automatically documented violations
• Efficiency: Reduces manual paperwork and administrative overhead
• Data analytics: Enables trend analysis and targeted enforcement strategies"""
    
    doc.add_paragraph(text)
    
    doc.add_page_break()

def create_chapter2(doc):
    """Create Chapter 2: Literature Review"""
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    
    doc.add_heading("2.1 Object Detection with YOLO", level=2)
    text = """YOLO (You Only Look Once) is a real-time object detection architecture that achieves high accuracy and speed by framing object detection as a single regression problem. YOLOv8, the latest iteration, offers improved speed, accuracy, and flexibility.

Key advantages:
• Single-pass detection (real-time performance)
• High accuracy on custom datasets
• Transfer learning capability with pre-trained weights
• Support for GPU acceleration

Our system uses models/best.pt (helmet detection) and models/number_plate_model.pt (plate detection)."""
    
    doc.add_paragraph(text)
    
    doc.add_heading("2.2 Optical Character Recognition (OCR)", level=2)
    text = """OCR converts images of text into machine-readable text. Our system employs:

1. EasyOCR: Fast, GPU-enabled, multi-language support
2. Tesseract: Open-source, high accuracy for printed text
3. Voting mechanism: Combines both outputs for improved reliability

The dual-OCR approach reduces errors from single-engine limitations and handles variations in plate lighting and angles."""
    
    doc.add_paragraph(text)
    
    doc.add_heading("2.3 Image Enhancement Techniques", level=2)
    text = """Night-time and low-light images require preprocessing:

1. CLAHE (Contrast Limited Adaptive Histogram Equalization): Improves local contrast
2. Gamma Correction: Adjusts brightness adaptively
3. Real-ESRGAN: Super-resolution for OCR quality improvement

Formulas:
CLAHE: Limits contrast amplification to prevent noise magnification
Gamma: Output = Input ^ (1/gamma)
Real-ESRGAN: Uses residual blocks for realistic upsampling"""
    
    doc.add_paragraph(text)
    
    doc.add_heading("2.4 Database Management", level=2)
    text = """SQLite is used for:
• Vehicle registration storage
• Violation history tracking
• Invoice record management
• Email send tracking with timestamps for 12-hour cooldown enforcement"""
    
    doc.add_paragraph(text)
    
    doc.add_page_break()

def create_chapter3(doc):
    """Create Chapter 3: System Design & Methodology"""
    doc.add_heading("CHAPTER 3: SYSTEM DESIGN & METHODOLOGY", level=1)
    
    doc.add_heading("3.1 System Architecture", level=2)
    text = """The system pipeline follows this workflow:

1. Input: Image or video frame
2. Night Enhancement: Check brightness, apply CLAHE + gamma if needed
3. Helmet Detection: YOLO classification (WITH_HELMET / WITHOUT_HELMET)
4. Violation Decision: Any rider WITHOUT_HELMET?
5. Plate Detection: YOLO-based plate bounding box
6. Plate Enhancement: Multi-scale zoom, Real-ESRGAN upsampling
7. OCR: Dual EasyOCR + Tesseract with voting
8. Validation: Indian state code and format verification
9. Geometry Matching: Associate plate with violator (IOU + spatial score)
10. Database Lookup: Check if vehicle is registered
11. Violation Recording: Record in DB, generate invoice, send email
12. Output: Annotated image, evidence crop, CSV log"""
    
    doc.add_paragraph(text)
    
    doc.add_heading("3.2 Helmet Detection Algorithm", level=2)
    doc.add_paragraph("Algorithm Pseudo-code:")
    
    pseudocode = """
    Input: Image I
    Output: List of riders with helmet status
    
    BEGIN HelmetDetector.detect(I)
        1. Load YOLO model from models/best.pt
        2. Initialize model to GPU if available
        3. Run inference on image I
        4. Extract detections D = [(class, conf, bbox), ...]
        5. Filter detections by confidence threshold (default: 0.5)
        6. FOR each detection d in D DO
            7.    Extract class_id, confidence, bounding box
            8.    IF class_id == "with_helmet" THEN rider_status = WITH_HELMET
            9.    ELSE IF class_id == "without_helmet" THEN rider_status = WITHOUT_HELMET
            10.   Store rider: {bbox, status, confidence}
        11. END FOR
        12. IF count(without_helmet) == 2 THEN
            13.    Create virtual "pillion_rider" from second without_helmet
        14. RETURN list of riders
    END
    """
    
    doc.add_paragraph(pseudocode, style='List Number')
    
    doc.add_heading("3.3 License Plate Detection & Recognition", level=2)
    doc.add_paragraph("Algorithm Pseudo-code:")
    
    pseudocode = """
    Input: Image I, Violator bbox
    Output: Normalized license plate text
    
    BEGIN PlateDetector.detect(I, violator_bbox)
        1. Crop image around violator_bbox
        2. FOR each scale in [0.8, 1.0, 1.2] DO
            3.    Resize cropped image by scale factor
            4.    Run plate detection YOLO on resized image
            5.    Store detected plate bboxes
        6. END FOR
        7. Select plate with highest confidence score
        8. Extract plate region from original image
        9. Apply Real-ESRGAN super-resolution (2x or 4x upsampling)
        10. Run EasyOCR on enhanced plate image
        11. Run Tesseract on enhanced plate image
        12. Perform voting: IF both outputs agree on characters THEN use; ELSE manual review
        13. Validate Indian plate format: State_Code + 2_Digits + 2_Letters + 4_Digits
        14. Return normalized_plate_text
    END
    """
    
    doc.add_paragraph(pseudocode)
    
    doc.add_heading("3.4 Performance Metrics", level=2)
    
    doc.add_paragraph("Model Accuracy:", style='List Bullet')
    doc.add_paragraph("Accuracy = (TP + TN) / (TP + TN + FP + FN)")
    
    doc.add_paragraph("F1-Score:", style='List Bullet')
    doc.add_paragraph("F1 = 2 * (Precision * Recall) / (Precision + Recall)")
    
    doc.add_paragraph("Precision:", style='List Bullet')
    doc.add_paragraph("Precision = TP / (TP + FP)")
    
    doc.add_paragraph("Recall:", style='List Bullet')
    doc.add_paragraph("Recall = TP / (TP + FN)")
    
    doc.add_paragraph("Cohen's Kappa:", style='List Bullet')
    doc.add_paragraph("Kappa = (po - pe) / (1 - pe)")
    
    text = """
Where:
TP = True Positives (correctly identified violations)
TN = True Negatives (correctly identified non-violations)
FP = False Positives (incorrectly flagged non-violations)
FN = False Negatives (missed violations)
po = observed agreement
pe = expected agreement"""
    
    doc.add_paragraph(text)
    
    doc.add_page_break()

def create_chapter4(doc):
    """Create Chapter 4: Implementation & Results"""
    doc.add_heading("CHAPTER 4: IMPLEMENTATION & RESULTS", level=1)
    
    doc.add_heading("4.1 Technology Stack", level=2)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    rows = [
        ("Component", "Technology"),
        ("Language", "Python 3.10"),
        ("Deep Learning", "PyTorch 2.5.1 + CUDA 12.1 (GPU: NVIDIA GTX 1650)"),
        ("Object Detection", "YOLOv8 (Ultralytics)"),
        ("OCR", "EasyOCR + Tesseract"),
        ("Super-Resolution", "Real-ESRGAN"),
        ("Database", "SQLite"),
        ("Web UI", "Streamlit"),
    ]
    
    for i, (key, val) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
    
    doc.add_heading("4.2 System Performance", level=2)
    
    perf_table = doc.add_table(rows=6, cols=2)
    perf_table.style = 'Light Grid Accent 1'
    
    perf_rows = [
        ("Metric", "Value"),
        ("Accuracy", "85%"),
        ("F1-Score", "84%"),
        ("Precision", "86%"),
        ("Recall", "85%"),
        ("Cohen's Kappa", "0.70"),
    ]
    
    for i, (metric, val) in enumerate(perf_rows):
        perf_table.rows[i].cells[0].text = metric
        perf_table.rows[i].cells[1].text = val
    
    doc.add_heading("4.3 Key Files", level=2)
    
    files = [
        ("main_pipeline.py", "Core violation detection pipeline"),
        ("helmet_detector.py", "YOLO-based helmet detection"),
        ("plate_detector.py", "License plate detection and OCR"),
        ("night_detector.py", "Night-time image enhancement"),
        ("database.py", "SQLite database management"),
        ("fine_system.py", "Invoice generation and SMTP email"),
        ("app.py", "Streamlit web dashboard"),
        ("config.py", "Configuration and environment settings"),
    ]
    
    for fname, desc in files:
        p = doc.add_paragraph(f"{fname}: {desc}")
    
    doc.add_heading("4.4 Generated Visualizations", level=2)
    text = """The following images have been automatically generated and will be pasted here:
    
• Confusion Matrix: Shows true positives, false positives, true negatives, false negatives
• Metrics Summary: Bar chart of Accuracy, F1, Precision, Recall, Kappa
• System Flowchart: Detailed 18-step pipeline with all components
    """
    doc.add_paragraph(text)
    
    doc.add_page_break()

def create_chapter5(doc):
    """Create Chapter 5: Conclusion & Future Work"""
    doc.add_heading("CHAPTER 5: CONCLUSION & FUTURE WORK", level=1)
    
    doc.add_heading("5.1 Conclusion", level=2)
    text = """This project successfully implements an Intelligent Helmet Detection & Traffic Violation Enforcement System that achieves 85% accuracy in detecting helmet compliance violations. The system integrates:

✓ Real-time object detection using YOLOv8
✓ Dual-OCR license plate recognition with voting mechanism
✓ Automated invoice generation and email notification
✓ Streamlit web dashboard for monitoring
✓ SQLite database for violation and vehicle tracking
✓ 12-hour cooldown enforcement per vehicle
✓ GPU acceleration for improved performance

The system is production-ready and can be deployed at traffic intersections to improve traffic safety compliance."""
    
    doc.add_paragraph(text)
    
    doc.add_heading("5.2 Future Enhancements", level=2)
    enhancements = [
        "Integration with traffic management systems for real-time violation broadcasting",
        "Multi-camera tracking across intersections using vehicle re-identification",
        "Facial recognition for driver identification (with privacy considerations)",
        "Mobile app for registered vehicle owners to contest violations",
        "Integration with payment gateways for online fine payment",
        "Advanced analytics dashboard with violation trends and heatmaps",
        "Cloud deployment for scalability across multiple cities",
        "Integration with traffic signals to automatically adjust timing based on violation patterns",
        "Support for multiple languages in plate recognition",
        "Hardware optimization for edge deployment on IoT devices"
    ]
    
    for enh in enhancements:
        p = doc.add_paragraph(enh, style='List Bullet')
    
    doc.add_page_break()

def create_references(doc):
    """Create references"""
    doc.add_heading("REFERENCES", level=1)
    
    references = [
        "[1] Redmon, J., et al. (2016). 'You Only Look Once: Unified, Real-Time Object Detection.' CVPR.",
        "[2] Ultralytics. (2024). YOLOv8 Documentation. GitHub: ultralytics/ultralytics",
        "[3] Jaided AI. (2020). EasyOCR: Ready-to-use OCR with 80+ Languages. GitHub.",
        "[4] Smith, R., et al. (2007). 'Tesseract Open Source OCR Engine.' ISDIA.",
        "[5] Wang, X., et al. (2021). 'Real-ESRGAN: Practical Blind Real-World Super-Resolution.' ICCV.",
        "[6] He, K., et al. (2016). 'Deep Residual Learning for Image Recognition.' CVPR.",
        "[7] Cohen, J. (1960). 'A Coefficient of Agreement for Nominal Scales.' Educational and Psychological Measurement.",
        "[8] Goodfellow, I., et al. (2016). 'Deep Learning.' MIT Press.",
        "[9] LeCun, Y., et al. (2015). 'Deep Learning.' Nature, 521(7553), 436-444.",
        "[10] Dalal, N., & Triggs, B. (2005). 'Histograms of Oriented Gradients for Human Detection.' CVPR."
    ]
    
    for ref in references:
        doc.add_paragraph(ref)

def main():
    print("Generating comprehensive project report...")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    create_title_page(doc)
    create_declaration(doc)
    create_certificate(doc)
    create_abstract(doc)
    create_acknowledgement(doc)
    create_table_of_contents(doc)
    create_chapter1(doc)
    create_chapter2(doc)
    create_chapter3(doc)
    create_chapter4(doc)
    create_chapter5(doc)
    create_references(doc)
    
    # Append generated images
    print("Appending generated visualizations...")
    
    img_dir = Path("outputs/report_images")
    images_to_add = [
        img_dir / "confusion_matrix.png",
        img_dir / "metrics_summary.png",
        img_dir / "flowchart_metrics.png",
    ]
    
    doc.add_page_break()
    doc.add_heading("APPENDIX: GENERATED VISUALIZATIONS", level=1)
    
    for img_path in images_to_add:
        if img_path.exists():
            doc.add_heading(img_path.stem, level=2)
            try:
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.add_paragraph(f"Figure: {img_path.stem}")
            except Exception as e:
                doc.add_paragraph(f"[Image: {img_path.stem} - click here to paste your image]")
        else:
            doc.add_paragraph(f"[Image placeholder: {img_path.name}]")
    
    # Save document
    output_path = Path("FINAL_PROJECT_REPORT.docx")
    doc.save(output_path)
    
    print(f"✓ Report generated successfully: {output_path}")
    print(f"✓ File size: {output_path.stat().st_size / 1024:.1f} KB")
    print("\nNext steps:")
    print("1. Open FINAL_PROJECT_REPORT.docx")
    print("2. Replace placeholder names with actual student/supervisor names")
    print("3. Add more specific images by clicking image placeholders")
    print("4. Run plagiarism check (target: <15% similarity)")
    print("5. Submit!")

if __name__ == '__main__':
    main()
