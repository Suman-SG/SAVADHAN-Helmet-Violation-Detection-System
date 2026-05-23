import os
import sys
import subprocess
import importlib
from pathlib import Path
import json

REQS = [
    ("pandas", "pandas"),
    ("numpy", "numpy"),
    ("sklearn", "scikit-learn"),
    ("matplotlib", "matplotlib"),
    ("seaborn", "seaborn"),
    ("docx", "python-docx"),
]


def ensure_packages():
    for mod_name, pip_name in REQS:
        try:
            importlib.import_module(mod_name)
        except ImportError:
            print(f"Installing {pip_name}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])


def find_label_prediction_sources(root=Path(".")):
    """Search workspace for CSV or JSON containing y_true / y_pred lists."""
    from pathlib import Path
    import pandas as pd

    root = Path(root)
    # Look for CSVs first
    for csv in root.rglob("*.csv"):
        try:
            df = pd.read_csv(csv)
            cols = [c.lower() for c in df.columns]
            if 'y_true' in cols and 'y_pred' in cols:
                return ('csv', csv, 'y_true', 'y_pred')
            # common alternative names
            if 'true' in cols and 'pred' in cols:
                return ('csv', csv, df.columns[cols.index('true')], df.columns[cols.index('pred')])
        except Exception:
            continue

    # Look for JSON containing arrays
    for j in root.rglob("*.json"):
        try:
            data = json.loads(j.read_text(encoding='utf-8'))
            if isinstance(data, dict) and 'y_true' in data and 'y_pred' in data:
                return ('json', j)
        except Exception:
            continue

    return (None, None)


def compute_and_plot(root=Path('.')):
    import pandas as pd
    import numpy as np
    import matplotlib.pyplot as plt
    import seaborn as sns
    from sklearn.metrics import (
        accuracy_score,
        f1_score,
        precision_score,
        recall_score,
        confusion_matrix,
        cohen_kappa_score,
        classification_report,
    )

    root = Path(root)
    out_dir = root / 'outputs' / 'report_images'
    out_dir.mkdir(parents=True, exist_ok=True)

    src = find_label_prediction_sources(root)
    results = {}

    if src[0] == 'csv':
        _, path, true_col, pred_col = src
        df = pd.read_csv(path)
        y_true = df[true_col]
        y_pred = df[pred_col]
    elif src[0] == 'json':
        _, path = src
        j = json.loads(Path(path).read_text(encoding='utf-8'))
        y_true = j['y_true']
        y_pred = j['y_pred']
    else:
        # try to extract metrics from outputs/metrics JSON files
        metrics_dir = root / 'outputs' / 'metrics'
        if metrics_dir.exists():
            for jf in metrics_dir.glob('*.json'):
                try:
                    data = json.loads(jf.read_text(encoding='utf-8'))
                    # try to take scalar metrics
                    for k in ('accuracy', 'f1_score', 'precision', 'recall', 'kappa'):
                        if k in data and k not in results:
                            results[k] = data[k]
                except Exception:
                    continue
        # Save a placeholder metrics image and return
        if results:
            fig, ax = plt.subplots(figsize=(6, 3))
            keys = list(results.keys())
            vals = [results[k] for k in keys]
            sns.barplot(x=vals, y=keys, palette='magma', ax=ax)
            ax.set_xlim(0, 1)
            ax.set_title('Available metrics from JSON')
            p = out_dir / 'metrics_available.png'
            fig.tight_layout()
            fig.savefig(p)
            plt.close(fig)
            return {'images': [str(p)], 'report': results}
        else:
            raise FileNotFoundError('No y_true/y_pred source found and no metrics JSON available.')

    # compute metrics
    labels = sorted(list(set(list(y_true) + list(y_pred))))
    acc = accuracy_score(y_true, y_pred)
    f1 = f1_score(y_true, y_pred, average='weighted')
    prec = precision_score(y_true, y_pred, average='weighted', zero_division=0)
    rec = recall_score(y_true, y_pred, average='weighted', zero_division=0)
    kappa = cohen_kappa_score(y_true, y_pred)
    cm = confusion_matrix(y_true, y_pred, labels=labels)
    clf_report = classification_report(y_true, y_pred, zero_division=0)

    results.update({'accuracy': acc, 'f1_score': f1, 'precision': prec, 'recall': rec, 'kappa': kappa, 'labels': labels})

    images = []

    # Confusion matrix heatmap
    fig, ax = plt.subplots(figsize=(6 + max(0, len(labels)-3), 5))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=labels, yticklabels=labels, ax=ax)
    ax.set_xlabel('Predicted')
    ax.set_ylabel('Actual')
    ax.set_title('Confusion Matrix')
    p1 = out_dir / 'confusion_matrix.png'
    fig.tight_layout()
    fig.savefig(p1)
    plt.close(fig)
    images.append(str(p1))

    # Metrics bar
    fig, ax = plt.subplots(figsize=(6, 3))
    keys = ['accuracy', 'f1_score', 'precision', 'recall', 'kappa']
    vals = [results[k] for k in keys]
    sns.barplot(x=vals, y=keys, palette='viridis', ax=ax)
    ax.set_xlim(0, 1)
    ax.set_title('Model summary metrics')
    p2 = out_dir / 'metrics_summary.png'
    fig.tight_layout()
    fig.savefig(p2)
    plt.close(fig)
    images.append(str(p2))

    # ROC/PR not added because labels may be multiclass; skip for now

    # Save textual report
    rep_path = out_dir / 'classification_report.txt'
    rep_path.write_text(clf_report, encoding='utf-8')

    # draw a simple flowchart annotated with key metrics
    fc_path = out_dir / 'flowchart_metrics.png'
    draw_flowchart_with_metrics(results, fc_path)
    images.append(str(fc_path))

    return {'images': images, 'report': results, 'classification_report': str(rep_path)}


def draw_flowchart_with_metrics(metrics, out_path):
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle

    fig, ax = plt.subplots(figsize=(14, 18))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 20)
    ax.axis('off')

    def box(x, y, w, h, text, color='#e6f2fb', edge='#2b7fb3', fontsize=9):
        """Draw a rounded box with text"""
        b = FancyBboxPatch((x-w/2, y-h/2), w, h, boxstyle='round,pad=0.1', linewidth=1.5, edgecolor=edge, facecolor=color)
        ax.add_patch(b)
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize, weight='bold', wrap=True)

    def arrow(x1, y1, x2, y2, label=''):
        """Draw an arrow with optional label"""
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle='->', mutation_scale=20, lw=2, color='#333'))
        if label:
            mid_x, mid_y = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mid_x + 0.3, mid_y, label, fontsize=8, style='italic', color='#666')

    # Main pipeline flow
    y_pos = 19
    
    # 1. Input
    box(5, y_pos, 2, 0.8, 'Input Image\n(JPEG/PNG)', color='#fff0cc')
    y_pos -= 1.2
    arrow(5, 19.6, 5, y_pos + 0.6)

    # 2. Night Detector
    box(5, y_pos, 3, 1.2, 'NightDetector.enhance()\n• Check brightness & std dev\n• Apply CLAHE + gamma', color='#ffe6e6')
    y_pos -= 1.8
    arrow(5, y_pos + 2.4, 5, y_pos + 0.6)

    # 3. Helmet Detector
    box(5, y_pos, 3.5, 1.5, 'HelmetDetector.detect()\n• YOLO (models/best.pt)\n• Classify: WITH/WITHOUT_HELMET\n• Create pillion riders', color='#e6f2fb')
    y_pos -= 2.1
    arrow(5, y_pos + 2.7, 5, y_pos + 0.9)

    # 4. Violation Check
    box(5, y_pos, 2.5, 1, 'Violation?\n(Any WITHOUT_HELMET)', color='#fff7cc')
    y_pos -= 1.8
    arrow(5, y_pos + 1.4, 5, y_pos + 0.75, 'YES')

    # 5. Plate Detector (left side detail)
    box(2, y_pos, 3.2, 1.8, 'PlateDetector.detect()\n• YOLO plate detection\n• Multi-scale zoom\n• Real-ESRGAN super-res\n• Dual OCR: EasyOCR+Tesseract', color='#f0e6ff')
    arrow(3.4, y_pos, 4, y_pos)
    
    # 6. Plate Validation (right side detail)
    box(8, y_pos, 3, 1.5, 'Plate Validation\n• Indian state code check\n• Format validation\n• FULL/PARTIAL/INVALID\n• Normalize plate text', color='#f0e6ff')
    arrow(6.5, y_pos, 6.8, y_pos)
    
    y_pos -= 2.2
    arrow(2, y_pos + 2.2, 3.5, y_pos + 0.8)
    arrow(8, y_pos + 2.2, 6.5, y_pos + 0.8)
    
    # 7. Geometry Matching
    box(5, y_pos, 3, 1, 'Geometry Matching\n(Match plate → violator)', color='#e6ffe6')
    y_pos -= 1.6
    arrow(5, y_pos + 1.1, 5, y_pos + 0.55)

    # 8. Save Evidence
    box(5, y_pos, 3, 0.9, 'Save Evidence Image\n(Rider + Plate crop)', color='#e6f0ff')
    y_pos -= 1.4
    arrow(5, y_pos + 0.95, 5, y_pos + 0.55)

    # 9. Database Lookup
    box(5, y_pos, 3, 1, 'Database Lookup\n(TrafficDatabase)', color='#ffe6f0')
    y_pos -= 1.6
    arrow(5, y_pos + 1.1, 5, y_pos + 0.65, 'Registered?')

    # 10a. If Registered (left)
    box(2.5, y_pos, 2.8, 1.5, 'If Registered:\n• Record violation\n• Generate PDF invoice\n• Send email (SMTP)\n• Log to CSV', color='#ccffcc')
    arrow(3.5, y_pos, 4, y_pos)
    
    # 10b. If Test Mode (right)
    box(7.5, y_pos, 2.8, 1.5, 'If Test Mode:\n• Send demo email\n  to configured\n  address\n• Log event', color='#ffcccc')
    arrow(6.5, y_pos, 6, y_pos)

    y_pos -= 2.1
    arrow(2.5, y_pos + 2.1, 4, y_pos + 0.7)
    arrow(7.5, y_pos + 2.1, 6, y_pos + 0.7)
    
    # 11. Final Output
    box(5, y_pos, 4, 1.2, 'Final Output:\nAnnotated Image | Evidence | Invoice |\nCSV Log | Email Sent', color='#99ff99', edge='#00aa00')

    # Metrics box
    metrics_txt = f"Performance Metrics\nAccuracy: {metrics.get('accuracy', 'N/A'):.1%}\nF1-Score: {metrics.get('f1_score', 'N/A'):.1%}\nPrecision: {metrics.get('precision', 'N/A'):.1%}\nRecall: {metrics.get('recall', 'N/A'):.1%}\nKappa: {metrics.get('kappa', 'N/A'):.3f}"
    ax.text(0.5, 1.5, metrics_txt, fontsize=8, bbox=dict(boxstyle='round', facecolor='#fff7e6', edgecolor='#d6a22a', pad=0.5), family='monospace')

    fig.savefig(out_path, bbox_inches='tight', dpi=150)
    plt.close(fig)


def append_to_docx(image_paths, docx_path="APROJECTREPORT.docx"):
    from docx import Document
    from docx.shared import Inches

    doc = Document(docx_path) if Path(docx_path).exists() else Document()
    doc.add_heading('Automated Analysis Images', level=2)

    for p in image_paths:
        fname = os.path.relpath(p)
        doc.add_paragraph(fname)
        try:
            doc.add_picture(p, width=Inches(5))
        except Exception:
            doc.add_paragraph('(image insertion failed)')

    doc.save(docx_path)


def main():
    ensure_packages()
    try:
        out = compute_and_plot(Path('.'))
        images = out.get('images', [])
        if images:
            append_to_docx(images)
        print('Analysis complete. Images saved to outputs/report_images/')
    except Exception as e:
        print('Error during analysis:', e)


if __name__ == '__main__':
    main()
