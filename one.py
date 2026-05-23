import os
import sys
import subprocess
import importlib
import glob
import json
from pathlib import Path

REQS = [
    ("requests", "requests"),
    ("matplotlib", "matplotlib"),
    ("seaborn", "seaborn"),
    ("docx", "python-docx"),
    ("PIL", "Pillow"),
    ("numpy", "numpy"),
]


def ensure_packages():
    for mod_name, pip_name in REQS:
        try:
            importlib.import_module(mod_name)
        except ImportError:
            print(f"Installing {pip_name}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])


def load_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def make_metrics_plots(out_dir):
    import matplotlib.pyplot as plt
    import seaborn as sns
    import numpy as np

    metrics_dir = Path("outputs") / "metrics"
    saved = []
    if not metrics_dir.exists():
        print("No outputs/metrics directory found; skipping metrics plots.")
        return saved

    for jfile in metrics_dir.glob("*.json"):
        data = load_json(jfile)
        if not data:
            continue

        name = jfile.stem
        out_path = Path(out_dir) / f"metrics_{name}.png"

        # Try to extract simple numeric dict for a barplot
        if isinstance(data, dict):
            nums = {k: v for k, v in data.items() if isinstance(v, (int, float))}
            if nums:
                sns.set(style="whitegrid")
                plt.figure(figsize=(8, 4))
                keys = list(nums.keys())
                vals = [nums[k] for k in keys]
                sns.barplot(x=vals, y=keys, palette="viridis")
                plt.title(f"Metrics: {name}")
                plt.tight_layout()
                plt.savefig(out_path)
                plt.close()
                saved.append(str(out_path))
                continue

        # If list of numbers or dicts, make a simple line plot
        if isinstance(data, list):
            try:
                arr = np.array(data)
                if arr.ndim == 1 and arr.size > 1:
                    plt.figure(figsize=(6, 3))
                    plt.plot(arr)
                    plt.title(f"Metrics Series: {name}")
                    plt.tight_layout()
                    plt.savefig(out_path)
                    plt.close()
                    saved.append(str(out_path))
                    continue
            except Exception:
                pass

    return saved


def make_placeholders(out_dir):
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle

    files = []
    placeholders = [
        ("method_flowchart.png", "Method Flowchart (paste real image here)"),
        ("system_architecture.png", "System Architecture Diagram (paste real image here)"),
        ("metrics_summary.png", "Metrics Summary Placeholder")
    ]

    for fname, text in placeholders:
        path = Path(out_dir) / fname
        fig, ax = plt.subplots(figsize=(6, 3.5))
        ax.add_patch(Rectangle((0.05, 0.05), 0.9, 0.9, fill=False, linewidth=2, edgecolor="#2b7fb3"))
        ax.text(0.5, 0.5, text, ha="center", va="center", fontsize=12)
        ax.axis("off")
        plt.tight_layout()
        fig.savefig(path)
        plt.close(fig)
        files.append(str(path))

    return files


def download_manifest(out_dir, manifest_path="images_to_download.txt"):
    import requests
    files = []
    mpath = Path(manifest_path)
    if not mpath.exists():
        print(f"No {manifest_path} found; skipping downloads.")
        return files

    with open(mpath, "r", encoding="utf-8") as f:
        for line in f:
            url = line.strip()
            if not url:
                continue
            try:
                r = requests.get(url, timeout=30)
                r.raise_for_status()
                name = os.path.basename(url.split("?")[0]) or f"img_{len(files)}.png"
                out_path = Path(out_dir) / name
                with open(out_path, "wb") as wf:
                    wf.write(r.content)
                files.append(str(out_path))
                print(f"Downloaded {url} -> {out_path}")
            except Exception as e:
                print(f"Failed to download {url}: {e}")

    return files


def append_to_docx(image_paths, docx_path="APROJECTREPORT.docx"):
    from docx import Document
    from docx.shared import Inches

    doc = Document(docx_path) if Path(docx_path).exists() else Document()
    doc.add_heading('Generated report images', level=2)

    for p in image_paths:
        fname = os.path.relpath(p)
        doc.add_paragraph(fname)
        try:
            doc.add_picture(p, width=Inches(5))
        except Exception:
            doc.add_paragraph('(thumbnail insertion failed)')

    doc.save(docx_path)
    print(f"Updated {docx_path} with {len(image_paths)} images.")


def main():
    ensure_packages()

    out_dir = Path("outputs") / "report_images"
    out_dir.mkdir(parents=True, exist_ok=True)

    created = []
    created += make_metrics_plots(out_dir)
    created += make_placeholders(out_dir)
    created += download_manifest(out_dir)

    # Deduplicate and keep relative paths
    created = [str(Path(p)) for p in dict.fromkeys(created)]

    if not created:
        print("No images were created or downloaded.")
    else:
        print("Images created/downloaded:")
        for c in created:
            print(" -", c)

    append_to_docx(created)

    print("Done. Open APROJECTREPORT.docx and paste/replace placeholders as needed.")


if __name__ == '__main__':
    main()
